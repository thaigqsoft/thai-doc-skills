#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
thai_docx.py — สร้างไฟล์ Word (.docx) ภาษาไทยให้เขียนเต็มบรรทัด จัดหน้าสวย
และรองรับรูปแบบเอกสารราชการ (ระเบียบงานสารบรรณ)

หลักการ:
  1) แทรก Zero-Width Space (U+200B) ระหว่างคำไทย เพื่อบอก Word ว่าตัดบรรทัดได้ตรงไหน
     → ข้อความไทยจึงเต็มบรรทัด ไม่ตัดก่อนเวลา
  2) ตั้งฟอนต์ให้ครบทั้งฝั่งอังกฤษ (ascii/hAnsi) และฝั่ง Complex Script (cs)
     ซึ่งเป็นฝั่งที่ภาษาไทยใช้จริง — จุดนี้ python-docx ปกติ "ไม่ตั้งให้"
     ทำให้ฟอนต์ไทยไม่เปลี่ยนตามที่สั่ง สคริปต์นี้แก้ให้แล้ว

รองรับ: หัวเรื่อง/หัวข้อย่อย, ย่อหน้า, bullet/เลขข้อ, ตาราง, รูปภาพ+คำบรรยาย,
สารบัญอัตโนมัติ, หัวกระดาษ/ท้ายกระดาษ, เลขหน้า, ตัวหนาแบบ **markdown**

ผู้เขียน: เผยแพร่เป็นสาธารณะ (MIT License) — ใช้ฟรี แก้ไข แจกจ่ายได้
"""

from __future__ import annotations

import re
import sys
import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from pythainlp import word_tokenize
    _HAS_PYTHAINLP = True
except Exception:  # pragma: no cover - pythainlp อาจยังไม่ได้ติดตั้ง
    _HAS_PYTHAINLP = False

    def word_tokenize(text, engine="newmm"):  # type: ignore
        # fallback แบบหยาบ: ถ้าไม่มี pythainlp ก็ไม่ตัดคำ (ยังใช้งานได้แต่บรรทัดอาจไม่เต็ม)
        return [text]


ZWS = "​"  # Zero-Width Space

# ─────────────────────────────────────────────────────────────────────────────
# ฟอนต์แห่งชาติที่มากับรีโป (โฟลเดอร์ fonts/) — key สั้น → ชื่อ family ที่ Word ใช้จริง
# ชื่อทางขวาต้องตรงเป๊ะตามที่ฝังในไฟล์ฟอนต์ ไม่งั้น Word จะหาไม่เจอ
# ─────────────────────────────────────────────────────────────────────────────
BUNDLED_FONTS = {
    "sarabun": "TH SarabunPSK",      # มาตรฐานเอกสารราชการ
    "krub": "TH Krub",
    "koho": "TH KoHo",
    "niramit": "TH Niramit AS",
    "kodchasal": "TH Kodchasal",
    "baijam": "TH Baijam",
    "chakrapetch": "TH Chakra Petch",
    "fahkwang": "TH Fah kwang",
    "k2d": "TH K2D July8",
    "mali": "TH Mali Grade 6",
}

# ชื่อยอดนิยมที่มักสะกดต่างกัน → ชื่อจริง
_FONT_ALIASES = {
    "th sarabun new": "TH Sarabun New",      # เวอร์ชันที่มากับ MS Office
    "sarabun new": "TH Sarabun New",
    "th sarabunpsk": "TH SarabunPSK",
    "sarabun psk": "TH SarabunPSK",
    "th sarabun psk": "TH SarabunPSK",
}


def resolve_font(name: str) -> str:
    """แปลงชื่อย่อ/ชื่อเล่นของฟอนต์เป็นชื่อ family จริงที่ Word ใช้
    รับได้ทั้ง key ('krub'), ชื่อเต็ม ('TH Krub'), หรือชื่อทั่วไป ('TH Sarabun New')"""
    if not name:
        return name
    low = name.strip().lower()
    if low in BUNDLED_FONTS:
        return BUNDLED_FONTS[low]
    if low in _FONT_ALIASES:
        return _FONT_ALIASES[low]
    return name


def list_bundled_fonts():
    """คืน dict ของฟอนต์ที่มากับรีโป (key → ชื่อ family)"""
    return dict(BUNDLED_FONTS)

# ─────────────────────────────────────────────────────────────────────────────
# Preset การจัดหน้า
# ─────────────────────────────────────────────────────────────────────────────
PRESETS = {
    # เอกสารราชการตามแนวระเบียบงานสารบรรณ: TH Sarabun New 16pt, A4
    # ขอบ: บน 2.5 / ล่าง 2.0 / ซ้าย 3.0 / ขวา 2.0 ซม. (ปรับได้)
    "saraban": dict(
        font_name="TH Sarabun New",
        font_size=16,
        page_size="A4",
        margins=dict(top=2.5, bottom=2.0, left=3.0, right=2.0),
        line_spacing=1.0,
        first_line_indent=1.25,
        space_after=0,
    ),
    # ทั่วไป/รายงาน: อ่านสบาย เว้นบรรทัด 1.5
    "default": dict(
        font_name="TH Sarabun New",
        font_size=14,
        page_size="A4",
        margins=dict(top=2.54, bottom=2.54, left=2.54, right=2.54),
        line_spacing=1.5,
        first_line_indent=0.0,
        space_after=6,
    ),
    # บทความ/หนังสือ: ย่อหน้าแบบมี indent ไม่เว้นบรรทัดระหว่างย่อหน้า
    "book": dict(
        font_name="TH Sarabun New",
        font_size=16,
        page_size="A4",
        margins=dict(top=2.54, bottom=2.54, left=3.0, right=2.54),
        line_spacing=1.3,
        first_line_indent=1.25,
        space_after=0,
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# การตัดคำ / แทรก ZWS
# ─────────────────────────────────────────────────────────────────────────────
_THAI_RUN = re.compile(r"([฀-๿]+)")


def insert_zwsp(text: str, engine: str = "newmm") -> str:
    """แทรก Zero-Width Space ระหว่างคำไทย — ภาษาอังกฤษ ตัวเลข URL ไม่ถูกแตะ"""
    if not text:
        return text
    parts = _THAI_RUN.split(text)
    out = []
    for part in parts:
        if part and _THAI_RUN.fullmatch(part):
            out.append(ZWS.join(word_tokenize(part, engine=engine)))
        else:
            out.append(part)
    return "".join(out)


# ─────────────────────────────────────────────────────────────────────────────
# ตัวช่วยระดับ XML — ตั้งฟอนต์ให้ครบทุกฝั่ง (รวม Complex Script ของไทย)
# ─────────────────────────────────────────────────────────────────────────────
def set_run_font(run, name=None, size=None, bold=None, italic=None, color=None):
    """
    ตั้งฟอนต์ของ run ให้ครบทั้ง ascii / hAnsi / cs (และ eastAsia)
    จุดสำคัญ: ฝั่ง cs (Complex Script) คือฝั่งที่อักษรไทยใช้จริง
    ถ้าตั้งแต่ run.font.name เฉย ๆ ฟอนต์ไทยจะไม่เปลี่ยน
    """
    rPr = run._element.get_or_add_rPr()

    if name:
        run.font.name = name  # ตั้ง ascii/hAnsi ให้
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), name)

    if size is not None:
        run.font.size = Pt(size)  # ตั้ง w:sz
        # ตั้ง w:szCs (ขนาดฝั่ง complex script) ให้ตรงกัน
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size * 2)))  # หน่วยเป็นครึ่ง pt

    if bold is not None:
        run.font.bold = bold
        # ตัวหนาฝั่ง complex script
        bCs = rPr.find(qn("w:bCs"))
        if bCs is None:
            bCs = OxmlElement("w:bCs")
            rPr.append(bCs)
        bCs.set(qn("w:val"), "true" if bold else "false")

    if italic is not None:
        run.font.italic = italic
        iCs = rPr.find(qn("w:iCs"))
        if iCs is None:
            iCs = OxmlElement("w:iCs")
            rPr.append(iCs)
        iCs.set(qn("w:val"), "true" if italic else "false")

    if color is not None:
        run.font.color.rgb = color


# ─────────────────────────────────────────────────────────────────────────────
# Inline markdown: รองรับ **ตัวหนา** และ *ตัวเอียง*
# ─────────────────────────────────────────────────────────────────────────────
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_formatted_runs(paragraph, text, font_name, font_size):
    """แตกข้อความออกเป็น run ตาม **ตัวหนา** / *ตัวเอียง* แล้วแทรก ZWS + ตั้งฟอนต์"""
    # tokenize เป็นลำดับ (style, ข้อความ)
    segments = [("normal", text)]

    def _split(segs, pattern, style):
        new = []
        for st, tx in segs:
            if st != "normal":
                new.append((st, tx))
                continue
            last = 0
            for m in pattern.finditer(tx):
                if m.start() > last:
                    new.append(("normal", tx[last:m.start()]))
                new.append((style, m.group(1)))
                last = m.end()
            if last < len(tx):
                new.append(("normal", tx[last:]))
        return new

    segments = _split(segments, _BOLD, "bold")
    segments = _split(segments, _ITALIC, "italic")

    if not any(tx for _, tx in segments):
        segments = [("normal", text)]

    for style, tx in segments:
        if tx == "":
            continue
        run = paragraph.add_run(insert_zwsp(tx))
        set_run_font(
            run,
            name=font_name,
            size=font_size,
            bold=(style == "bold") or None,
            italic=(style == "italic") or None,
        )


# ─────────────────────────────────────────────────────────────────────────────
# Field helpers (เลขหน้า / สารบัญ)
# ─────────────────────────────────────────────────────────────────────────────
def _add_field(paragraph, instr, font_name=None, font_size=None):
    """เพิ่ม field code (เช่น PAGE, NUMPAGES, TOC) ลงในย่อหน้า"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    if font_name or font_size:
        set_run_font(run, name=font_name, size=font_size)
    return run


def _add_page_number(paragraph, fmt, font_name, font_size, align):
    """เลขหน้า; fmt รองรับ {n}=เลขหน้าปัจจุบัน {total}=จำนวนหน้ารวม"""
    paragraph.alignment = align
    # แตก fmt ออกตาม {n} / {total}
    tokens = re.split(r"(\{n\}|\{total\})", fmt)
    for tok in tokens:
        if tok == "{n}":
            _add_field(paragraph, "PAGE", font_name, font_size)
        elif tok == "{total}":
            _add_field(paragraph, "NUMPAGES", font_name, font_size)
        elif tok:
            run = paragraph.add_run(insert_zwsp(tok))
            set_run_font(run, name=font_name, size=font_size)


def _add_toc(doc, font_name, font_size, title="สารบัญ"):
    """แทรกสารบัญอัตโนมัติ (Word จะอัปเดตเลขหน้าเมื่อกด Update Field / F9)"""
    if title:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(h.add_run(title), name=font_name, size=font_size + 4, bold=True)

    p = doc.add_paragraph()
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u', font_name, font_size)
    # ข้อความชั่วคราวให้ผู้ใช้รู้ว่าต้องอัปเดต field
    note = doc.add_paragraph()
    set_run_font(
        note.add_run("(คลิกขวาที่สารบัญ → Update Field เพื่อให้เลขหน้าแสดงผล)"),
        name=font_name,
        size=font_size - 2,
        italic=True,
        color=RGBColor(0x80, 0x80, 0x80),
    )


# ─────────────────────────────────────────────────────────────────────────────
# Section / page setup
# ─────────────────────────────────────────────────────────────────────────────
def _apply_page_setup(section, page_size, margins):
    if page_size == "A4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    elif page_size == "Letter":
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
    m = margins or {}
    section.top_margin = Cm(m.get("top", 2.54))
    section.bottom_margin = Cm(m.get("bottom", 2.54))
    section.left_margin = Cm(m.get("left", 2.54))
    section.right_margin = Cm(m.get("right", 2.54))


def _set_styles(doc, font_name, font_size):
    """ตั้งฟอนต์ default ของ Normal + Heading ให้ครบฝั่ง cs"""
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    # ตั้ง cs ให้ style Normal ผ่าน XML
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font_name)

    for level, size in [(1, font_size + 6), (2, font_size + 4), (3, font_size + 2)]:
        try:
            hstyle = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        hstyle.font.name = font_name
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        hr = hstyle.element.get_or_add_rPr()
        hrf = hr.find(qn("w:rFonts"))
        if hrf is None:
            hrf = OxmlElement("w:rFonts")
            hr.append(hrf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            hrf.set(qn(attr), font_name)


# ─────────────────────────────────────────────────────────────────────────────
# ตัวสร้างเนื้อหาแต่ละชนิด
# ─────────────────────────────────────────────────────────────────────────────
def _add_table(doc, item, font_name, font_size):
    rows = item["rows"]
    if not rows:
        return
    has_header = item.get("header", True)
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = item.get("style", "Table Grid")
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""  # เคลียร์ค่าเริ่มต้น
            val = row[j] if j < len(row) else ""
            run = cell.paragraphs[0].add_run(insert_zwsp(str(val)))
            set_run_font(
                run, name=font_name, size=font_size,
                bold=(has_header and i == 0) or None,
            )
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if (has_header and i == 0) else WD_ALIGN_PARAGRAPH.LEFT


def _add_image(doc, item, font_name, font_size):
    from docx.shared import Inches
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = item.get("width")
    if width:
        run.add_picture(item["path"], width=Cm(width))
    else:
        run.add_picture(item["path"])
    cap = item.get("caption")
    if cap:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cp.add_run(insert_zwsp(cap)), name=font_name, size=font_size - 1, italic=True)


# ─────────────────────────────────────────────────────────────────────────────
# ฟังก์ชันหลัก
# ─────────────────────────────────────────────────────────────────────────────
def create_docx(
    paragraphs: list,
    output_path: str,
    *,
    preset: str = None,
    font_name: str = None,
    font_size: int = None,
    page_size: str = None,
    margins: dict = None,
    line_spacing: float = None,
    first_line_indent: float = None,
    space_after: int = None,
    header_text: str = None,
    footer_text: str = None,
    page_number: str = None,          # "footer-center" | "footer-right" | "header-right" | None
    page_number_format: str = "{n}",  # เช่น "หน้า {n}" หรือ "{n} / {total}"
    toc: bool = False,
    title: str = None,
    embed_fonts: bool = True,         # ฝังฟอนต์ลงไฟล์เป็นค่าเริ่มต้น (เปิดเครื่องไหนก็เห็นฟอนต์ถูก)
):
    """
    สร้าง Word document จากรายการ paragraph

    แต่ละ item เป็น dict:
        {"text": "...", "type": "body"}                     # ย่อหน้า (รองรับ **ตัวหนา**)
        {"text": "...", "type": "title"}                    # ชื่อเรื่อง (กลางหน้า ตัวใหญ่)
        {"text": "...", "type": "subtitle"}
        {"text": "...", "type": "heading1"|"heading2"|"heading3"}
        {"text": "...", "type": "bullet", "level": 0}       # หัวข้อจุด
        {"text": "...", "type": "number", "level": 0}       # หัวข้อเลข
        {"text": "...", "type": "quote"}
        {"text": "...", "type": "caption"}
        {"type": "table", "rows": [[...],[...]], "header": True}
        {"type": "image", "path": "a.png", "width": 12, "caption": "ภาพที่ 1"}
        {"type": "pagebreak"}

    preset: "saraban" (ราชการ) | "default" | "book" — ตั้งค่าเริ่มต้นการจัดหน้า
            พารามิเตอร์ที่ระบุเองจะทับค่าของ preset
    """
    cfg = dict(PRESETS.get(preset, PRESETS["default"]))
    # override ด้วยค่าที่ผู้ใช้ระบุ
    if font_name is not None: cfg["font_name"] = font_name
    if font_size is not None: cfg["font_size"] = font_size
    if page_size is not None: cfg["page_size"] = page_size
    if margins is not None: cfg["margins"] = margins
    if line_spacing is not None: cfg["line_spacing"] = line_spacing
    if first_line_indent is not None: cfg["first_line_indent"] = first_line_indent
    if space_after is not None: cfg["space_after"] = space_after

    font_name = resolve_font(cfg["font_name"])
    font_size = cfg["font_size"]

    doc = Document()
    _set_styles(doc, font_name, font_size)

    section = doc.sections[0]
    _apply_page_setup(section, cfg["page_size"], cfg["margins"])

    # หัว/ท้ายกระดาษ + เลขหน้า
    if header_text:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(hp.add_run(insert_zwsp(header_text)), name=font_name, size=font_size - 2)
    if footer_text:
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(fp.add_run(insert_zwsp(footer_text)), name=font_name, size=font_size - 2)
    if page_number:
        loc, _, side = page_number.partition("-")
        align = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
        }.get(side, WD_ALIGN_PARAGRAPH.CENTER)
        target = section.footer if loc == "footer" else section.header
        # ถ้ามีข้อความ footer/header อยู่แล้ว ให้สร้างย่อหน้าใหม่
        pgp = target.paragraphs[0] if not (target.paragraphs[0].text) else target.add_paragraph()
        _add_page_number(pgp, page_number_format, font_name, font_size - 2, align)

    # ชื่อเรื่องหลัก (ถ้าส่ง title มา)
    if title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(tp.add_run(insert_zwsp(title)), name=font_name, size=font_size + 8, bold=True)
        tp.paragraph_format.space_after = Pt(12)

    # สารบัญ
    if toc:
        _add_toc(doc, font_name, font_size)
        doc.add_page_break()

    ls = cfg["line_spacing"]
    indent = cfg["first_line_indent"]
    sp_after = cfg["space_after"]

    for item in paragraphs:
        ptype = item.get("type", "body")

        if ptype == "pagebreak":
            doc.add_page_break()
            continue
        if ptype == "table":
            _add_table(doc, item, font_name, font_size)
            doc.add_paragraph().paragraph_format.space_after = Pt(sp_after)
            continue
        if ptype == "image":
            _add_image(doc, item, font_name, font_size)
            continue

        text = item.get("text", "")

        if ptype == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(insert_zwsp(text)), name=font_name, size=font_size + 8, bold=True)
            p.paragraph_format.space_after = Pt(12)

        elif ptype == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(insert_zwsp(text)), name=font_name, size=font_size + 2, italic=True)
            p.paragraph_format.space_after = Pt(10)

        elif ptype.startswith("heading"):
            lvl = ptype.replace("heading", "")
            level = int(lvl) if lvl.isdigit() else 1
            level = max(1, min(level, 3))
            p = doc.add_heading("", level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_formatted_runs(p, text, font_name, font_size + (8 - 2 * level))
            for r in p.runs:
                r.font.bold = True

        elif ptype in ("bullet", "number"):
            style = "List Bullet" if ptype == "bullet" else "List Number"
            lvl = int(item.get("level", 0))
            sname = style if lvl == 0 else f"{style} {min(lvl + 1, 3)}"
            try:
                p = doc.add_paragraph(style=sname)
            except KeyError:
                p = doc.add_paragraph(style=style)
            p.paragraph_format.line_spacing = ls
            _add_formatted_runs(p, text, font_name, font_size)

        elif ptype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.line_spacing = ls
            _add_formatted_runs(p, text, font_name, font_size)
            for r in p.runs:
                set_run_font(r, italic=True, color=RGBColor(0x55, 0x55, 0x55))

        elif ptype == "caption":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_runs(p, text, font_name, font_size - 1)
            for r in p.runs:
                set_run_font(r, italic=True)

        else:  # body
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = ls
            pf.space_after = Pt(sp_after)
            if indent:
                pf.first_line_indent = Cm(indent)
            _add_formatted_runs(p, text, font_name, font_size)

    doc.save(output_path)

    # ฝังฟอนต์ลงไฟล์ (default) — ให้พกพาไปเครื่องที่ไม่มีฟอนต์ไทยได้
    if embed_fonts:
        try:
            from font_embed import embed_fonts as _embed
        except ImportError:
            from .font_embed import embed_fonts as _embed  # type: ignore
        try:
            done = _embed(output_path, [font_name])
            if not done:
                print(f"⚠️  ฝังฟอนต์ไม่สำเร็จ: ไม่พบไฟล์ฟอนต์ของ '{font_name}' "
                      f"(เอกสารยังใช้ได้ แต่ต้องลงฟอนต์ที่เครื่องปลายทาง)")
        except Exception as e:  # ฝังพลาดไม่ควรทำให้สร้างเอกสารล้มเหลว
            print(f"⚠️  ข้ามการฝังฟอนต์ ({e})")

    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# แปลง markdown/ข้อความธรรมดา → รายการ paragraph
# ─────────────────────────────────────────────────────────────────────────────
def parse_markdown(raw: str) -> list:
    """แปลงข้อความแบบ markdown ง่าย ๆ เป็นรายการ paragraph สำหรับ create_docx()"""
    items = []
    lines = raw.replace("\r\n", "\n").split("\n")
    buf = []

    def flush():
        if buf:
            items.append({"text": " ".join(buf).strip(), "type": "body"})
            buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush()
            i += 1
            continue

        # หัวข้อแบบ # / ## / ###
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            flush()
            level = len(m.group(1))
            items.append({"text": m.group(2), "type": f"heading{level}"})
            i += 1
            continue

        # bullet: - หรือ *
        if re.match(r"^[-*]\s+", stripped):
            flush()
            items.append({"text": re.sub(r"^[-*]\s+", "", stripped), "type": "bullet"})
            i += 1
            continue

        # เลขข้อ: 1. 2) ...
        if re.match(r"^\d+[.)]\s+", stripped):
            flush()
            items.append({"text": re.sub(r"^\d+[.)]\s+", "", stripped), "type": "number"})
            i += 1
            continue

        # หัวข้อเลขที่แบบราชการ: 1.2.3 ...
        if re.match(r"^\d+(\.\d+)*\s+\S", stripped) and len(stripped) < 100:
            flush()
            depth = stripped.split()[0].count(".")
            items.append({"text": stripped, "type": f"heading{min(depth + 1, 3)}"})
            i += 1
            continue

        buf.append(stripped)
        i += 1

    flush()
    return items


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────
def _safe_stdout():
    """ทำให้ stdout/stderr พิมพ์ภาษาไทย/emoji ได้บนคอนโซล Windows (cp874) โดยไม่ crash"""
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


def main():
    _safe_stdout()
    parser = argparse.ArgumentParser(
        description="สร้าง Word (.docx) ภาษาไทยเต็มบรรทัด รองรับเอกสารราชการ",
    )
    parser.add_argument("input", nargs="?", help="ไฟล์ข้อความ/markdown (.txt, .md)")
    parser.add_argument("-o", "--output", default="output.docx", help="ไฟล์ผลลัพธ์ (.docx)")
    parser.add_argument("--preset", default="default", choices=list(PRESETS), help="รูปแบบการจัดหน้า")
    parser.add_argument("--font", default=None,
                        help="ฟอนต์ (ทับค่า preset) — ใส่ชื่อเต็มหรือ key เช่น krub, koho, sarabun")
    parser.add_argument("--list-fonts", action="store_true", help="แสดงฟอนต์ที่มากับรีโปแล้วออก")
    parser.add_argument("--size", type=int, default=None, help="ขนาดฟอนต์ pt (ทับค่า preset)")
    parser.add_argument("--title", default=None, help="ชื่อเรื่องบนหัวเอกสาร")
    parser.add_argument("--toc", action="store_true", help="ใส่สารบัญอัตโนมัติ")
    parser.add_argument("--page-number", default=None,
                        help="ตำแหน่งเลขหน้า เช่น footer-center, footer-right")
    parser.add_argument("--header", default=None, help="ข้อความหัวกระดาษ")
    parser.add_argument("--footer", default=None, help="ข้อความท้ายกระดาษ")
    parser.add_argument("--no-embed", action="store_true",
                        help="ไม่ฝังฟอนต์ลงไฟล์ (ค่าเริ่มต้นคือฝัง)")
    args = parser.parse_args()

    if args.list_fonts:
        print("ฟอนต์ที่มากับรีโป (ใช้เป็น --font <key>):")
        for k, v in list_bundled_fonts().items():
            print(f"  {k:14s} → {v}")
        return

    if not args.input:
        parser.error("ต้องระบุไฟล์ input (หรือใช้ --list-fonts)")

    with open(args.input, "r", encoding="utf-8") as f:
        raw = f.read()

    paragraphs = parse_markdown(raw)
    create_docx(
        paragraphs,
        args.output,
        preset=args.preset,
        font_name=args.font,
        font_size=args.size,
        title=args.title,
        toc=args.toc,
        page_number=args.page_number,
        header_text=args.header,
        footer_text=args.footer,
        embed_fonts=not args.no_embed,
    )
    if not _HAS_PYTHAINLP:
        print("⚠️  ไม่พบ pythainlp — ข้อความไทยอาจไม่เต็มบรรทัด (pip install pythainlp)")
    print(f"✅ สร้างไฟล์: {args.output}")


if __name__ == "__main__":
    main()
